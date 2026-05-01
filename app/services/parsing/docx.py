"""Word .docx reader.

Walks paragraphs in document order and reconstructs a markdown-ish text
stream so that downstream `MarkdownNodeParser` can recover heading
hierarchy. Heading level n → `"#" * n` prefix.
"""
from __future__ import annotations

import re
from pathlib import Path

import docx as python_docx

from app.services.parsing.base import Document, make_base_metadata

_HEADING_LEVEL_RE = re.compile(r"^Heading\s+(\d+)$", re.IGNORECASE)


class DocxReader:
    def load_data(self, file_path: str | Path) -> list[Document]:
        path = Path(file_path)
        doc = python_docx.Document(str(path))

        lines: list[str] = []
        breadcrumb_stack: list[str] = []
        # We'll emit ONE Document holding the entire file (in markdown form).
        # Heading-level chunking is done by SentenceSplitter+MarkdownNodeParser
        # during the chunking layer.

        for p in doc.paragraphs:
            txt = (p.text or "").strip()
            if not txt:
                continue

            level = _heading_level(p.style.name if p.style else None)
            if level:
                # Trim/extend breadcrumb stack to this level
                breadcrumb_stack = breadcrumb_stack[: level - 1] + [txt]
                lines.append(f"{'#' * level} {txt}")
            else:
                lines.append(txt)

        # Tables → render as markdown tables
        for tbl in doc.tables:
            rendered = _table_to_markdown(tbl)
            if rendered:
                lines.append(rendered)

        if not lines:
            return []

        full_text = "\n\n".join(lines)
        md = make_base_metadata(
            file_path=path,
            parser="docx",
            page=None,
            section_idx=0,
        )
        return [Document(text=full_text, metadata=md)]


def _heading_level(style_name: str | None) -> int | None:
    if not style_name:
        return None
    m = _HEADING_LEVEL_RE.match(style_name)
    if m:
        n = int(m.group(1))
        return min(n, 6)
    if style_name.lower() == "title":
        return 1
    return None


def _table_to_markdown(tbl) -> str:
    rows: list[list[str]] = []
    for row in tbl.rows:
        rows.append([(c.text or "").replace("\n", " ").strip() for c in row.cells])
    if not rows:
        return ""
    header = rows[0]
    sep = ["---"] * len(header)
    body = rows[1:]
    out = ["| " + " | ".join(header) + " |",
           "| " + " | ".join(sep) + " |"]
    for r in body:
        out.append("| " + " | ".join(r) + " |")
    return "\n".join(out)
